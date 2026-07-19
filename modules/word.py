import docx
import os

def extract_text_from_docx(docx_path):
    """
    從 DOCX 檔案中提取所有文字內容。
    """
    try:
        doc = docx.Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)
    except Exception as e:
        return f"Error extracting text from DOCX: {e}"

def get_docx_metadata(docx_path):
    """
    獲取 DOCX 檔案的元數據。
    """
    try:
        doc = docx.Document(docx_path)
        prop = doc.core_properties
        metadata = {
            "Title": prop.title,
            "Author": prop.author,
            "Category": prop.category,
            "Comments": prop.comments,
            "Content Status": prop.content_status,
            "Created": prop.created,
            "Identifier": prop.identifier,
            "Keywords": prop.keywords,
            "Language": prop.language,
            "Last Modified By": prop.last_modified_by,
            "Last Printed": prop.last_printed,
            "Modified": prop.modified,
            "Revision": prop.revision,
            "Subject": prop.subject,
            "Version": prop.version
        }
        # 過濾掉 None 值
        return {k: v for k, v in metadata.items() if v is not None}
    except Exception as e:
        return f"Error getting DOCX metadata: {e}"

if __name__ == '__main__':
    # 創建一個簡單的測試 DOCX 檔案
    test_doc_path = "test_document.docx"
    document = docx.Document()
    document.add_heading('測試文件標題', level=1)
    document.add_paragraph('這是第一段文字。')
    document.add_paragraph('這是第二段文字，包含一些')
    document.add_paragraph('關鍵字：Python, DOCX, 測試。')
    document.save(test_doc_path)

    print(f"Created test DOCX file: {test_doc_path}")

    # 測試文字提取
    extracted_text = extract_text_from_docx(test_doc_path)
    print("\nExtracted Text from DOCX:")
    print(extracted_text)

    # 測試元數據獲取
    metadata = get_docx_metadata(test_doc_path)
    print("\nDOCX Metadata:")
    for key, value in metadata.items():
        print(f"  {key}: {value}")

    # 清理測試檔案
    os.remove(test_doc_path)
    print(f"\nRemoved test DOCX file: {test_doc_path}")
